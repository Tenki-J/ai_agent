from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\SBS\.claude\skills\docx-creator-workspace\iteration-1\eval-design-spec\without_skill\outputs\web_design_spec.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: set table cell background ─────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

# ── Helper: style table header row ────────────────────────────────────────────
def style_header_row(row, bg_hex="2E4057", fg_hex=(255, 255, 255)):
    for cell in row.cells:
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold       = True
                run.font.color.rgb = RGBColor(*fg_hex)
                run.font.size  = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Helper: zebra-stripe body rows ────────────────────────────────────────────
def style_body_rows(table, even_bg="EAF0FB", odd_bg="FFFFFF"):
    for i, row in enumerate(table.rows[1:], start=1):
        bg = even_bg if i % 2 == 0 else odd_bg
        for cell in row.cells:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ── Helper: add a colour swatch cell ──────────────────────────────────────────
def add_swatch_cell(row, hex_val: str):
    cell = row.cells[1]          # HEX column
    set_cell_bg(cell, hex_val)
    for para in cell.paragraphs:
        for run in para.runs:
            # choose black or white label depending on brightness
            r, g, b = int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16)
            luminance = 0.299*r + 0.587*g + 0.114*b
            run.font.color.rgb = RGBColor(0, 0, 0) if luminance > 128 else RGBColor(255, 255, 255)
            run.bold = True

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Web Design Specification", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

meta = doc.add_paragraph("Version 1.0  ·  2026-06-07  ·  백곰이 Auto-generated")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
meta.runs[0].font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — COLOUR PALETTE
# ═══════════════════════════════════════════════════════════════════════════════
h1 = doc.add_heading("1. Colour Palette", level=1)
h1.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph(
    "The following palette defines all official brand and UI colours. "
    "Use only these values to maintain visual consistency across every surface."
)

colour_data = [
    ("Primary Blue",        "1A73E8", "Main CTA buttons, links, active states"),
    ("Primary Dark",        "2E4057", "Navigation bar, hero headings"),
    ("Accent Teal",         "00BFA5", "Highlights, badges, success indicators"),
    ("Accent Amber",        "FFA000", "Warnings, promotional tags, star ratings"),
    ("Background Light",    "F8F9FA", "Page background, card surfaces"),
    ("Surface White",       "FFFFFF", "Modal dialogs, input fields, tooltips"),
    ("Text Primary",        "212121", "Body text, headings, labels"),
    ("Text Secondary",      "757575", "Captions, placeholders, metadata"),
    ("Border / Divider",    "E0E0E0", "Table borders, section dividers, input borders"),
    ("Error / Danger",      "D32F2F", "Error messages, destructive actions"),
    ("Success Green",       "388E3C", "Success toasts, confirmation states"),
    ("Disabled Grey",       "BDBDBD", "Disabled buttons, inactive controls"),
]

colour_table = doc.add_table(rows=1, cols=3)
colour_table.style = "Table Grid"
colour_table.alignment = WD_TABLE_ALIGNMENT.CENTER
colour_table.columns[0].width = Inches(1.8)
colour_table.columns[1].width = Inches(1.4)
colour_table.columns[2].width = Inches(3.6)

hdr = colour_table.rows[0]
for cell, text in zip(hdr.cells, ["Colour Name", "HEX Value", "Usage / Purpose"]):
    cell.text = text
style_header_row(hdr)

for name, hex_val, usage in colour_data:
    row = colour_table.add_row()
    row.cells[0].text = name
    row.cells[1].text = f"#{hex_val}"
    row.cells[2].text = usage
    add_swatch_cell(row, hex_val)

style_body_rows(colour_table)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TYPOGRAPHY
# ═══════════════════════════════════════════════════════════════════════════════
h2 = doc.add_heading("2. Typography Specifications", level=1)
h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph(
    "All type is set in the Inter typeface family (variable, Google Fonts). "
    "Sizes are in px assuming a 16 px root; rem equivalents are noted where relevant."
)

typo_data = [
    ("Display Heading",  "Inter",   "56 px / 3.5 rem",  "700 Bold",      "Landing-page hero titles"),
    ("H1 — Page Title",  "Inter",   "40 px / 2.5 rem",  "700 Bold",      "Top-level section headings"),
    ("H2 — Section",     "Inter",   "32 px / 2 rem",    "600 SemiBold",  "Card titles, modal headings"),
    ("H3 — Sub-section", "Inter",   "24 px / 1.5 rem",  "600 SemiBold",  "Widget headings, sidebar titles"),
    ("H4 — Label",       "Inter",   "20 px / 1.25 rem", "500 Medium",    "Form group labels, accordion headers"),
    ("Body — Large",     "Inter",   "18 px / 1.125 rem","400 Regular",   "Intro paragraphs, feature descriptions"),
    ("Body — Default",   "Inter",   "16 px / 1 rem",    "400 Regular",   "Standard body copy, list items"),
    ("Body — Small",     "Inter",   "14 px / 0.875 rem","400 Regular",   "Secondary descriptions, table cells"),
    ("Caption",          "Inter",   "12 px / 0.75 rem", "400 Regular",   "Image captions, timestamps, tooltips"),
    ("Overline",         "Inter",   "11 px / 0.6875 rem","600 SemiBold", "Category tags, step labels (ALL CAPS)"),
    ("Button — Large",   "Inter",   "16 px / 1 rem",    "600 SemiBold",  "Primary & secondary large buttons"),
    ("Button — Default", "Inter",   "14 px / 0.875 rem","600 SemiBold",  "Standard action buttons"),
    ("Code / Mono",      "JetBrains Mono", "14 px / 0.875 rem","400 Regular","Inline code, code blocks, terminal"),
]

typo_table = doc.add_table(rows=1, cols=5)
typo_table.style = "Table Grid"
typo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
col_widths = [1.5, 1.4, 1.4, 1.3, 2.7]
for i, w in enumerate(col_widths):
    typo_table.columns[i].width = Inches(w)

typo_hdr = typo_table.rows[0]
for cell, text in zip(typo_hdr.cells, ["Element", "Font Family", "Size", "Weight", "Usage"]):
    cell.text = text
style_header_row(typo_hdr)

for row_data in typo_data:
    row = typo_table.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

style_body_rows(typo_table)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — LAYOUT RULES
# ═══════════════════════════════════════════════════════════════════════════════
h3 = doc.add_heading("3. Layout Rules", level=1)
h3.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

# ── 3a Spacing / Margin ───────────────────────────────────────────────────────
doc.add_heading("3.1  Spacing & Margin Scale", level=2)
doc.add_paragraph(
    "Spacing follows an 8 px base unit. All margins, paddings, and gaps must be "
    "multiples of this value to maintain visual rhythm."
)

spacing_data = [
    ("xs",  "4 px",   "0.25 rem", "Icon gap, tight inline spacing"),
    ("sm",  "8 px",   "0.5 rem",  "Input inner padding, list item gap"),
    ("md",  "16 px",  "1 rem",    "Card inner padding, form group gap"),
    ("lg",  "24 px",  "1.5 rem",  "Section top/bottom padding"),
    ("xl",  "32 px",  "2 rem",    "Component separation, modal padding"),
    ("2xl", "48 px",  "3 rem",    "Large section margins"),
    ("3xl", "64 px",  "4 rem",    "Hero vertical padding"),
    ("4xl", "96 px",  "6 rem",    "Page-level section dividers"),
    ("5xl", "128 px", "8 rem",    "Full-bleed hero section height minimum"),
]

spacing_table = doc.add_table(rows=1, cols=4)
spacing_table.style = "Table Grid"
spacing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([0.8, 0.9, 1.0, 5.6]):
    spacing_table.columns[i].width = Inches(w)

sp_hdr = spacing_table.rows[0]
for cell, text in zip(sp_hdr.cells, ["Token", "px", "rem", "Recommended Usage"]):
    cell.text = text
style_header_row(sp_hdr)

for row_data in spacing_data:
    row = spacing_table.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

style_body_rows(spacing_table)
doc.add_paragraph()

# ── 3b Grid ───────────────────────────────────────────────────────────────────
doc.add_heading("3.2  Grid System", level=2)
doc.add_paragraph(
    "The layout uses a 12-column CSS Grid. Column behaviour adapts per breakpoint. "
    "All max-widths are centred with auto horizontal margins."
)

grid_data = [
    ("Max Content Width", "1280 px",  "Constrains content on large screens"),
    ("Number of Columns", "12",       "Flexible column-span system"),
    ("Gutter (column gap)", "24 px",  "Space between adjacent columns"),
    ("Outer Margin — Desktop", "48 px", "Left & right page margin ≥ 1280 px"),
    ("Outer Margin — Tablet",  "32 px", "Left & right page margin 768–1279 px"),
    ("Outer Margin — Mobile",  "16 px", "Left & right page margin < 768 px"),
    ("Row Gap (default)",   "32 px",  "Vertical gap between grid rows"),
    ("Card Column Span",    "4 cols", "Standard 3-up card layout on desktop"),
    ("Sidebar Width",       "3 cols", "Right sidebar for article / dashboard pages"),
    ("Main Content Width",  "9 cols", "Content area when sidebar is present"),
    ("Full-width Section",  "12 cols","Hero, banner, footer — edge-to-edge"),
]

grid_table = doc.add_table(rows=1, cols=3)
grid_table.style = "Table Grid"
grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([2.2, 1.3, 4.8]):
    grid_table.columns[i].width = Inches(w)

g_hdr = grid_table.rows[0]
for cell, text in zip(g_hdr.cells, ["Property", "Value", "Notes"]):
    cell.text = text
style_header_row(g_hdr)

for row_data in grid_data:
    row = grid_table.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

style_body_rows(grid_table)
doc.add_paragraph()

# ── 3c Responsive Breakpoints ─────────────────────────────────────────────────
doc.add_heading("3.3  Responsive Breakpoints", level=2)
doc.add_paragraph(
    "Breakpoints follow a mobile-first strategy. CSS min-width media queries are used. "
    "The table below defines each tier, its target device class, column count, and key behaviour changes."
)

bp_data = [
    ("xs  — Extra Small", "0 px",     "< 480 px",   "4",  "Single-column stacking; hamburger nav; font scale –10%"),
    ("sm  — Small",       "480 px",   "< 768 px",   "4",  "Two-column grid for cards; bottom-tab navigation"),
    ("md  — Medium",      "768 px",   "< 1024 px",  "8",  "Tablet layout; collapsible sidebar; 8-col grid"),
    ("lg  — Large",       "1024 px",  "< 1280 px",  "12", "Full desktop nav; 3-col cards; sidebar visible"),
    ("xl  — Extra Large", "1280 px",  "< 1536 px",  "12", "Max content width active (1280 px); desktop optimised"),
    ("2xl — 2× Large",    "1536 px",  "≥ 1536 px",  "12", "Wide layout; optional 4-col card grid; larger hero"),
]

bp_table = doc.add_table(rows=1, cols=5)
bp_table.style = "Table Grid"
bp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([1.5, 0.9, 0.9, 0.6, 4.4]):
    bp_table.columns[i].width = Inches(w)

bp_hdr = bp_table.rows[0]
for cell, text in zip(bp_hdr.cells, ["Breakpoint", "min-width", "max-width", "Cols", "Layout Behaviour"]):
    cell.text = text
style_header_row(bp_hdr)

for row_data in bp_data:
    row = bp_table.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

style_body_rows(bp_table)
doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
note = doc.add_paragraph(
    "This document was auto-generated by 백곰이 (Claude Code Agent) on 2026-06-07. "
    "All values are subject to design review before production implementation."
)
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
