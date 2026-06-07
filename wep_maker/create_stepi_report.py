"""
STEPI 과학기술정책연구원 — 브랜드 CI 반영 보고서
Brand colors: STEPI Blue #003087 / STEPI Orange #E87722
Typography: 맑은 고딕
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors ────────────────────────────────────────────
STEPI_BLUE   = RGBColor(0x00, 0x30, 0x87)
STEPI_ORANGE = RGBColor(0xE8, 0x77, 0x22)
STEPI_NAVY   = RGBColor(0x00, 0x1B, 0x55)
GRAY_TEXT    = RGBColor(0x37, 0x41, 0x51)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_C = RGBColor(0xEB, 0xF3, 0xFB)

KO_FONT = "맑은 고딕"


def set_run_font(run, size_pt=None, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = KO_FONT
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), KO_FONT)
    rFonts.set(qn("w:hAnsi"), KO_FONT)
    rFonts.set(qn("w:eastAsia"), KO_FONT)
    rPr.insert(0, rFonts)
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, color_hex="E87722", sz="12"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_divider(doc, color="D1D5DB"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_banner(doc, label, title):
    """오렌지 배너 (섹션 헤더)"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False

    # 왼쪽 번호 칸 (네이비)
    lc = tbl.cell(0, 0)
    lc.width = Inches(0.45)
    set_cell_bg(lc, "001B55")
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(7)
    lp.paragraph_format.space_after = Pt(7)
    r = lp.add_run(label)
    set_run_font(r, 9, bold=True, color=STEPI_ORANGE)

    # 오른쪽 제목 칸 (STEPI 블루)
    rc = tbl.cell(0, 1)
    set_cell_bg(rc, "003087")
    rp = rc.paragraphs[0]
    rp.paragraph_format.left_indent = Cm(0.4)
    rp.paragraph_format.space_before = Pt(7)
    rp.paragraph_format.space_after = Pt(7)
    r2 = rp.add_run(title)
    set_run_font(r2, 13, bold=True, color=WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_body_para(doc, text, indent=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = Pt(19)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    set_run_font(run, 10.5, color=GRAY_TEXT)


# ════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.2)

# ── 헤더 ─────────────────────────────────────────────────
hp = section.header.paragraphs[0]
hp.clear()
r1 = hp.add_run("STEPI  과학기술정책연구원")
set_run_font(r1, 8, bold=True, color=STEPI_BLUE)
r2 = hp.add_run("     |     기관 개요 및 정책 연구 프로필")
set_run_font(r2, 8, color=GRAY_TEXT)
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ── 푸터 ─────────────────────────────────────────────────
fp = section.footer.paragraphs[0]
fp.clear()
rf = fp.add_run("www.stepi.re.kr     |     작성일: 2026-06-07     |     과학기술정책연구원")
set_run_font(rf, 8, color=GRAY_TEXT)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════
# 커버 블록
# ════════════════════════════════════════════════════════════
cover = doc.add_table(rows=1, cols=1)
cc = cover.cell(0, 0)
set_cell_bg(cc, "001B55")

cp = cc.paragraphs[0]
cp.paragraph_format.space_before = Pt(22)
cp.paragraph_format.space_after = Pt(4)
cp.paragraph_format.left_indent = Cm(0.6)

r_sub = cp.add_run("기관 개요 보고서\n")
set_run_font(r_sub, 10, italic=True, color=RGBColor(0xCC, 0xDD, 0xF0))
r_main = cp.add_run("과학기술정책연구원\n")
set_run_font(r_main, 26, bold=True, color=WHITE)
r_en = cp.add_run("Science and Technology Policy Institute (STEPI)")
set_run_font(r_en, 11, color=RGBColor(0xCC, 0xDD, 0xF0))

# 오렌지 액센트 바
orange = doc.add_table(rows=1, cols=1)
oc = orange.cell(0, 0)
set_cell_bg(oc, "E87722")
op = oc.paragraphs[0]
op.paragraph_format.space_before = Pt(4)
op.paragraph_format.space_after = Pt(4)
op.paragraph_format.left_indent = Cm(0.6)
r_meta = op.add_run("설립: 1987년   |   위치: 세종시   |   주무부처: 과학기술정보통신부")
set_run_font(r_meta, 9, bold=True, color=WHITE)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════
# 01 서론
# ════════════════════════════════════════════════════════════
add_section_banner(doc, "01", "서론")

add_body_para(doc,
    "과학기술정책연구원(STEPI)은 1987년 대한민국 정부의 출연으로 설립된 정부출연연구기관(GFRI)입니다. "
    "과학기술정보통신부(MSIT) 산하에서 운영되며 세종시에 본원을 두고 있습니다.")
add_body_para(doc,
    "STEPI의 법적 사명은 국가 과학기술혁신(STI) 정책에 관한 독립적이고 증거 기반의 연구를 수행하여 "
    "국가 거버넌스를 지원하는 것입니다. 한국의 장기 STI 전략을 설계·평가·개선하는 핵심 분석 기관으로 기능합니다.")
add_body_para(doc,
    "연구원은 세 가지 핵심 기능 축을 중심으로 운영됩니다: 정책 연구 및 분석, 통계 데이터 생산, "
    "국제 STI 협력. 이를 통해 과학기술기본계획(5개년), 과학기술기본법, 국가 R&D 예산 배분 과정에 직접적인 기여를 합니다.")

add_divider(doc)

# ════════════════════════════════════════════════════════════
# 02 핵심 연구 분야
# ════════════════════════════════════════════════════════════
add_section_banner(doc, "02", "핵심 연구 분야")

research_tracks = [
    ("국가혁신시스템(NIS) 분석",
     "산·학·연 협력 구조, 기술이전 메커니즘, 혁신 클러스터 역동성 등 한국의 국가 혁신 생태계를 체계적으로 분석하고 평가합니다."),
    ("R&D 투자 및 프로그램 평가",
     "국가 R&D 지출의 효율성과 사회적 수익률을 평가하고, 국가과학기술자문회의에 예산 배분 우선순위와 평가 프레임워크를 제안합니다."),
    ("기술예측 및 미래전망",
     "격년으로 실시하는 한국기술예측조사를 포함한 체계적인 기술예측 활동을 통해 10~20년 후 등장 기술과 그 사회경제적 함의를 분석합니다."),
    ("과학기술 인력 연구",
     "이공계 인력의 공급·수요·질적 수준을 분석하고, 대학원 교육 정책, 연구자 경력경로, 첨단 기술 분야 인력난 해소 전략을 연구합니다."),
    ("규제과학 및 기술 거버넌스",
     "인공지능·바이오기술·양자컴퓨팅 등 신흥기술에 관한 규제 프레임워크를 연구하고, 혁신 촉진과 공공 안전·윤리 간 균형을 위한 정책 방안을 제시합니다."),
    ("국제 STI 협력",
     "양자·다자 STI 파트너십을 분석하고, OECD 회원국 대비 한국의 혁신 수준을 벤치마킹하며, EU·미국·역내 주요국과의 국제 공동연구 협력 방안을 자문합니다."),
]

for i, (title, body) in enumerate(research_tracks):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False

    num_c = tbl.cell(0, 0)
    num_c.width = Inches(0.45)
    set_cell_bg(num_c, "003087")
    np2 = num_c.paragraphs[0]
    np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    np2.paragraph_format.space_before = Pt(6)
    np2.paragraph_format.space_after = Pt(6)
    rn = np2.add_run(str(i + 1).zfill(2))
    set_run_font(rn, 14, bold=True, color=WHITE)

    txt_c = tbl.cell(0, 1)
    set_cell_bg(txt_c, "EBF3FB")
    tp = txt_c.paragraphs[0]
    tp.paragraph_format.left_indent = Cm(0.3)
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(6)
    rt = tp.add_run(title + "\n")
    set_run_font(rt, 10.5, bold=True, color=STEPI_BLUE)
    rb = tp.add_run(body)
    set_run_font(rb, 9.5, color=GRAY_TEXT)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

add_divider(doc)

# ════════════════════════════════════════════════════════════
# 03 정책 기여
# ════════════════════════════════════════════════════════════
add_section_banner(doc, "03", "정책 기여")

pi = doc.add_paragraph(style="Normal")
pi.paragraph_format.space_after = Pt(8)
ri = pi.add_run(
    "STEPI의 연구 성과는 구체적인 정책 수단과 입법 프레임워크로 전환됩니다. "
    "다음은 가장 대표적인 정책 기여 사례입니다.")
set_run_font(ri, 10.5, color=GRAY_TEXT)

impact_items = [
    ("과학기술기본계획(5개년)",
     "STEPI는 역대 과학기술기본계획의 분석적 토대를 제공합니다. 진단 보고서, "
     "기술격차 분석, 이해관계자 조사 데이터가 각 계획의 전략적 우선순위에 직접 반영됩니다."),
    ("한국기업혁신조사(KIS)",
     "1996년 이후 격년으로 발간되는 KIS는 제조업·서비스업의 혁신 활동을 측정하는 핵심 통계 조사입니다. "
     "OECD 오슬로 매뉴얼을 준수하며 Eurostat 및 OECD에 국제 비교용 데이터를 제출합니다."),
    ("연구개발활동조사(R&D 통계연보)",
     "한국의 R&D 투자, 연구원 수, 특허 활동에 관한 권위 있는 연간 데이터를 편찬·발행하며, "
     "OECD 주요과학기술지표(MSTI)에서 공식 인용 출처로 활용됩니다."),
    ("기술수준평가",
     "분야별 전문가 패널과 협력하여 120개 이상의 전략 기술 분야에서 한국의 기술 수준을 "
     "미국·EU·일본·중국과 비교하며, 국가 R&D 사업 우선순위 설정에 직결됩니다."),
    ("AI 및 신흥기술 거버넌스 프레임워크",
     "한국의 국가 인공지능 전략(2019) 수립과 이후 AI법 논의에 기여하는 정책 제언서 및 입법 권고안을 작성하였으며, "
     "알고리즘 투명성·공공부문 AI 활용 조항에 직접 영향을 미쳤습니다."),
]

for i, (title, body) in enumerate(impact_items):
    pi2 = doc.add_paragraph(style="Normal")
    pi2.paragraph_format.space_before = Pt(6)
    pi2.paragraph_format.space_after = Pt(2)
    ro = pi2.add_run(f"  {i+1}  ")
    set_run_font(ro, 10.5, bold=True, color=STEPI_ORANGE)
    rt2 = pi2.add_run(title)
    set_run_font(rt2, 10.5, bold=True, color=STEPI_BLUE)

    pb = doc.add_paragraph(style="Normal")
    pb.paragraph_format.left_indent = Cm(0.8)
    pb.paragraph_format.space_after = Pt(4)
    pb.paragraph_format.line_spacing = Pt(18)
    rb2 = pb.add_run(body)
    set_run_font(rb2, 10, color=GRAY_TEXT)

    if i < len(impact_items) - 1:
        add_divider(doc, "E5E7EB")

add_divider(doc)

# ════════════════════════════════════════════════════════════
# 04 결론 (네이비 박스)
# ════════════════════════════════════════════════════════════
add_section_banner(doc, "04", "결론")

conclusion_paras = [
    "STEPI는 한국 과학기술 거버넌스 체계의 핵심 기관으로서, 약 40년에 걸쳐 급격한 산업 전환과 "
    "압축적 기술 추격이라는 국내 환경 속에서 증거 기반 STI 정책 수립의 관행을 제도화해 왔습니다.",
    "정책 연구 기관이자 통계 당국이라는 이중적 역할 덕분에, STEPI는 의제 설정·전략 설계에서부터 "
    "프로그램 평가와 법률 개정에 이르는 정책 전 주기에 걸쳐 독보적인 영향력을 행사합니다.",
    "인구 감소, 반도체 및 첨단 제조업에서의 지정학적 경쟁 심화, AI 전환이 요구하는 구조적 변화 등 "
    "복합적 도전에 대응하는 과정에서, STEPI의 연구 의제는 프론티어 기술 거버넌스, 공급망 회복력, "
    "신흥기술 국제 표준화 규범 형성을 향해 더욱 집중될 것으로 전망됩니다.",
]

conc_tbl = doc.add_table(rows=1, cols=1)
conc_c = conc_tbl.cell(0, 0)
set_cell_bg(conc_c, "001B55")
set_cell_border(conc_c, "E87722", "12")

for i, text in enumerate(conclusion_paras):
    cp3 = conc_c.paragraphs[0] if i == 0 else conc_c.add_paragraph()
    cp3.paragraph_format.left_indent = Cm(0.5)
    cp3.paragraph_format.right_indent = Cm(0.5)
    cp3.paragraph_format.space_before = Pt(8) if i == 0 else Pt(6)
    cp3.paragraph_format.space_after = Pt(6) if i < 2 else Pt(14)
    cp3.paragraph_format.line_spacing = Pt(19)
    rc3 = cp3.add_run(text)
    set_run_font(rc3, 10.5, color=WHITE)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

src_p = doc.add_paragraph()
src_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rs = src_p.add_run("출처: 과학기술정책연구원 (www.stepi.re.kr)")
set_run_font(rs, 8, italic=True, color=GRAY_TEXT)

# ── 저장 ──────────────────────────────────────────────────
OUTPUT = r"C:\Users\SBS\Desktop\agent ai-2\wep_maker\stepi_report.docx"
try:
    doc.save(OUTPUT)
    print(f"[OK] 저장 완료: {OUTPUT}")
except Exception as e:
    print(f"[ERROR] {e}")
