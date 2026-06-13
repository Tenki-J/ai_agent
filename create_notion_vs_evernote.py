from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\SBS\.claude\skills\docx-creator-workspace\iteration-1\eval-website-comparison\without_skill\outputs\notion_vs_evernote.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_bold(cell, bold=True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold

def style_header_row(row, bg_hex='2E4057'):
    for cell in row.cells:
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def main():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # =====================
    # TITLE
    # =====================
    title = doc.add_heading('Notion vs Evernote: 서비스 비교 분석', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
        run.font.size = Pt(20)

    doc.add_paragraph()

    intro = doc.add_paragraph(
        'Notion과 Evernote는 세계적으로 널리 사용되는 메모 및 생산성 도구입니다. '
        '두 서비스는 정보 관리와 노트 작성이라는 공통 목적을 가지지만, '
        '기능 철학, 협업 방식, 가격 정책 등 여러 면에서 뚜렷한 차이가 있습니다. '
        '이 문서에서는 주요 항목별로 두 서비스를 비교합니다.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # =====================
    # SECTION 1: 가격 비교
    # =====================
    add_heading(doc, '1. 가격 비교 (Pricing)', level=1)

    add_body(doc,
        'Notion과 Evernote 모두 무료 플랜을 제공하지만, 유료 플랜의 구조와 가격대에는 차이가 있습니다. '
        'Notion은 개인 사용자에게 무제한 페이지를 무료로 제공하며, 팀 협업 기능을 위한 Plus/Business 플랜을 운영합니다. '
        'Evernote는 무료 플랜의 기기 수 제한이 강하고, 유료 플랜(Personal/Professional)을 통해 전체 기능을 해제합니다.'
    )

    price_table = doc.add_table(rows=1, cols=3)
    price_table.style = 'Table Grid'
    price_table.autofit = False
    price_table.columns[0].width = Cm(4)
    price_table.columns[1].width = Cm(6)
    price_table.columns[2].width = Cm(6)

    hdr = price_table.rows[0]
    hdr.cells[0].text = '플랜 구분'
    hdr.cells[1].text = 'Notion'
    hdr.cells[2].text = 'Evernote'
    style_header_row(hdr)

    price_data = [
        ('무료 플랜', '무제한 페이지, 무제한 블록\n(게스트 1명 초대)', '2개 기기 한정, 60MB 업로드/월'),
        ('기본 유료\n(Personal)', '$10/월 (Plus)\n연간 결제 시 $8/월', '$14.99/월\n연간 결제 시 $10.83/월'),
        ('고급 유료\n(Business/Professional)', '$18/월 (Business)\n연간 결제 시 $15/월', '$17.99/월\n연간 결제 시 $14.17/월'),
        ('기업용 플랜', 'Enterprise (문의)', 'Teams (문의)'),
    ]

    for row_data in price_data:
        row = price_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # =====================
    # SECTION 2: 주요 기능 비교
    # =====================
    add_heading(doc, '2. 주요 기능 비교 (Core Features)', level=1)

    add_body(doc,
        'Notion은 데이터베이스, 캔버스 보드, 위키, 페이지 임베드 등 매우 유연한 구조로 복잡한 정보 체계를 구성할 수 있습니다. '
        '반면 Evernote는 태그 기반의 강력한 노트 분류 체계와 웹 클리퍼, PDF 주석 기능 등 수집 및 정리에 특화되어 있습니다. '
        'AI 기능 측면에서는 두 서비스 모두 점차 AI 통합을 강화하고 있습니다.'
    )

    feat_table = doc.add_table(rows=1, cols=3)
    feat_table.style = 'Table Grid'
    feat_table.autofit = False
    feat_table.columns[0].width = Cm(5)
    feat_table.columns[1].width = Cm(5.5)
    feat_table.columns[2].width = Cm(5.5)

    hdr2 = feat_table.rows[0]
    hdr2.cells[0].text = '기능 항목'
    hdr2.cells[1].text = 'Notion'
    hdr2.cells[2].text = 'Evernote'
    style_header_row(hdr2)

    feat_data = [
        ('노트 에디터', '블록 기반 에디터\n(텍스트, 이미지, 코드 등)', '리치 텍스트 에디터\n(첨부파일, 체크리스트)'),
        ('데이터베이스', '표, 캘린더, 보드, 갤러리, 타임라인\n다양한 뷰 지원', '미지원\n(테이블 기능 제한적)'),
        ('템플릿', '다양한 공식/커뮤니티 템플릿 제공', '기본 노트 템플릿 제공'),
        ('웹 클리퍼', '크롬 확장 프로그램 지원', '강력한 웹 클리퍼\n(여러 클립 형식 지원)'),
        ('AI 기능', 'Notion AI (유료 애드온)\n요약, 번역, 자동완성', 'Evernote AI (통합 중)\n노트 요약, 검색 보조'),
        ('오프라인 지원', '유료 플랜에서 지원', '유료 플랜에서 지원'),
        ('검색 기능', '전문 텍스트 검색 (유료 강화)', '이미지/PDF 내 텍스트 인식\n(OCR) 검색 지원'),
    ]

    for row_data in feat_data:
        row = feat_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # =====================
    # SECTION 3: 협업 기능 비교
    # =====================
    add_heading(doc, '3. 협업 기능 비교 (Collaboration)', level=1)

    add_body(doc,
        'Notion은 팀 협업을 핵심 가치로 설계된 도구로, 실시간 공동 편집, 멘션, 댓글, 권한 설정 등 협업 기능이 풍부합니다. '
        'Evernote는 원래 개인용 메모 도구로 출발해 협업 기능이 상대적으로 제한적이며, '
        '팀 공유는 가능하지만 Notion에 비해 실시간 협업 경험이 부족합니다.'
    )

    collab_table = doc.add_table(rows=1, cols=3)
    collab_table.style = 'Table Grid'
    collab_table.autofit = False
    collab_table.columns[0].width = Cm(5)
    collab_table.columns[1].width = Cm(5.5)
    collab_table.columns[2].width = Cm(5.5)

    hdr3 = collab_table.rows[0]
    hdr3.cells[0].text = '협업 항목'
    hdr3.cells[1].text = 'Notion'
    hdr3.cells[2].text = 'Evernote'
    style_header_row(hdr3)

    collab_data = [
        ('실시간 공동 편집', '지원 (여러 사용자 동시 편집)', '제한적 지원'),
        ('댓글 & 멘션', '페이지/블록 단위 댓글, @멘션 지원', '노트 댓글 지원\n(@멘션 기능 미흡)'),
        ('권한 관리', '페이지별 세밀한 권한 설정\n(편집, 보기, 댓글 등)', '노트/노트북 단위 공유\n(편집/보기 구분)'),
        ('게스트 초대', '무료: 게스트 1명\n유료: 다수 게스트', '유료 플랜에서 팀 공유'),
        ('팀 워크스페이스', '팀 전용 워크스페이스 제공', 'Spaces 기능 (Teams 플랜)'),
        ('알림 기능', '페이지 변경, 댓글 알림', '노트 공유 알림'),
    ]

    for row_data in collab_data:
        row = collab_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # =====================
    # SECTION 4: 플랫폼 지원 비교
    # =====================
    add_heading(doc, '4. 플랫폼 지원 비교 (Platform Support)', level=1)

    add_body(doc,
        '두 서비스 모두 주요 운영체제와 모바일 플랫폼을 지원하며 웹 브라우저를 통해서도 접근 가능합니다. '
        'Notion은 Linux 데스크톱 앱을 공식 지원하고, API를 통한 외부 연동이 활발합니다. '
        'Evernote는 역사적으로 다양한 플랫폼을 지원해왔으나 최근 일부 플랫폼 지원을 축소했습니다.'
    )

    platform_table = doc.add_table(rows=1, cols=3)
    platform_table.style = 'Table Grid'
    platform_table.autofit = False
    platform_table.columns[0].width = Cm(5)
    platform_table.columns[1].width = Cm(5.5)
    platform_table.columns[2].width = Cm(5.5)

    hdr4 = platform_table.rows[0]
    hdr4.cells[0].text = '플랫폼'
    hdr4.cells[1].text = 'Notion'
    hdr4.cells[2].text = 'Evernote'
    style_header_row(hdr4)

    platform_data = [
        ('웹 브라우저', '지원 (모든 주요 브라우저)', '지원 (모든 주요 브라우저)'),
        ('Windows', '데스크톱 앱 지원', '데스크톱 앱 지원'),
        ('macOS', '데스크톱 앱 지원', '데스크톱 앱 지원'),
        ('Linux', '공식 데스크톱 앱 지원', '미지원 (웹 브라우저 사용)'),
        ('iOS / iPadOS', '앱 지원', '앱 지원'),
        ('Android', '앱 지원', '앱 지원'),
        ('Apple Watch', '미지원', '미지원'),
        ('외부 연동 (API)', 'Notion API 공개\n(Zapier, Slack, GitHub 등)', 'IFTTT, Zapier 연동 지원\n(공식 API 제한적)'),
        ('이메일 전송 저장', '미지원', '이메일을 노트로 전송 가능'),
    ]

    for row_data in platform_data:
        row = platform_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # =====================
    # SECTION 5: 종합 요약
    # =====================
    add_heading(doc, '5. 종합 요약 및 추천', level=1)

    add_body(doc,
        '두 서비스는 각기 다른 사용자 층을 공략하고 있습니다. '
        'Notion은 팀 협업, 프로젝트 관리, 데이터베이스 기반 워크플로우를 필요로 하는 사용자에게 적합합니다. '
        '유연한 구조로 개인 위키부터 팀 운영 도구까지 폭넓게 활용할 수 있습니다. '
        'Evernote는 방대한 노트를 빠르게 캡처하고 체계적으로 정리하는 데 강점이 있으며, '
        'OCR 검색과 웹 클리퍼를 자주 사용하는 연구자나 콘텐츠 수집가에게 유리합니다.'
    )

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.autofit = False
    summary_table.columns[0].width = Cm(5)
    summary_table.columns[1].width = Cm(5.5)
    summary_table.columns[2].width = Cm(5.5)

    hdr5 = summary_table.rows[0]
    hdr5.cells[0].text = '항목'
    hdr5.cells[1].text = 'Notion'
    hdr5.cells[2].text = 'Evernote'
    style_header_row(hdr5)

    summary_data = [
        ('추천 대상', '팀 협업, 프로젝트 관리,\n정보 아키텍처 구성 필요 사용자', '개인 메모, 정보 수집,\nOCR 검색 활용 사용자'),
        ('학습 난이도', '중간~높음\n(기능이 많아 초기 진입장벽)', '낮음~중간\n(직관적인 UI)'),
        ('데이터 이식성', 'Markdown, CSV, HTML 내보내기', 'ENEX 형식 내보내기'),
        ('전반적 강점', '유연성, 협업, 구조화', '정보 수집, 검색, 안정성'),
    ]

    for row_data in summary_data:
        row = summary_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[i].paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    # =====================
    # FOOTER NOTE
    # =====================
    doc.add_paragraph()
    note = doc.add_paragraph('※ 본 문서의 가격 및 기능 정보는 2024년 기준이며, 실제 서비스 정책에 따라 변경될 수 있습니다.')
    note.paragraph_format.space_before = Pt(12)
    for run in note.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.italic = True

    # =====================
    # SAVE
    # =====================
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")
    print("Sections included:")
    print("  - Title & Introduction")
    print("  - Section 1: 가격 비교 (Pricing) — 4-row comparison table")
    print("  - Section 2: 주요 기능 비교 (Core Features) — 7-row comparison table")
    print("  - Section 3: 협업 기능 비교 (Collaboration) — 6-row comparison table")
    print("  - Section 4: 플랫폼 지원 비교 (Platform Support) — 9-row comparison table")
    print("  - Section 5: 종합 요약 및 추천 (Summary) — 4-row comparison table")
    print("  - Footer note")

if __name__ == '__main__':
    main()
