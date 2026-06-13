# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

FONT_NAME = '맑은 고딕'

doc = Document()

# 문서 기본 스타일(Normal)에 동아시아 폰트 설정
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)

# 페이지 여백
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)


def apply_font(run, name=FONT_NAME):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)


def set_heading(paragraph, text, level=2, color=RGBColor(0x8B, 0x00, 0x57)):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    apply_font(run)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after  = Pt(4)


def add_para(doc, text, size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    apply_font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    apply_font(run)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 표지 ──────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('S.E.S. 요약 보고서')
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(0x8B, 0x00, 0x57)
apply_font(title_run)

doc.add_paragraph()

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run('출처: 위키백과 한국어판 (ko.wikipedia.org/wiki/S.E.S.)')
sub_run.font.size = Pt(10)
sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
apply_font(sub_run)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('작성일: ' + datetime.date.today().strftime('%Y년 %m월 %d일'))
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
apply_font(date_run)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ── 1. 개요 ───────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '1. 개요 (Overview)', level=2)

add_para(doc,
    'S.E.S.(에스이에스)는 바다(Bada), 유진(Eugene), 슈(Shoo)로 구성된 대한민국의 3인조 걸그룹으로, '
    'SM엔터테인먼트 최초의 걸 그룹이다. 1997년 데뷔해 1990년대 후반~2000년대 초반 한국 가요계를 '
    '대표하는 걸그룹으로 자리매김했다. 청순한 이미지와 팝 스타일의 음악으로 대중의 큰 사랑을 받았으며, '
    '현재 전 세계를 사로잡고 있는 K-Pop 걸그룹들의 원조격으로 평가받는다. '
    '그룹명 S.E.S.는 세 멤버의 영문 이니셜(Sea, Eugene, Shoo)을 딴 것이다.'
)
doc.add_paragraph()

# ── 2. 멤버 ───────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '2. 멤버 소개', level=2)

members = [
    ('바다 (Bada / 최성희)',
     '뛰어난 음악 실력으로 그룹의 보컬을 담당. 강렬하고 파워풀한 가창력으로 팬들에게 '
     '깊은 인상을 남겼다. 해체 이후에도 솔로 가수로 활발히 활동하며 뮤지컬과 음악 활동을 이어가고 있다.'),
    ('유진 (Eugene / 김유진)',
     '뛰어난 외모로 그룹의 비주얼을 담당. 활동 당시 청순한 이미지로 많은 팬들의 사랑을 받았으며, '
     '해체 이후 배우로 전향해 드라마와 영화에서 활발히 활동하고 있다.'),
    ('슈 (Shoo / 오유수)',
     '특유의 매력으로 팬들을 사로잡는 역할을 담당. 밝고 사랑스러운 이미지로 그룹의 분위기를 '
     '이끌었으며, 해체 이후에도 대중적 인지도를 유지하고 있다.'),
]

for name, desc in members:
    add_para(doc, f'■ {name}', bold=True)
    add_para(doc, desc, indent=True)
doc.add_paragraph()

# ── 3. 데뷔 및 활동 역사 ─────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '3. 데뷔 및 활동 역사', level=2)

history = [
    ('① 결성 및 데뷔 (1996~1997)',
     '1996년 SM엔터테인먼트에서 결성, 1997년 데뷔 타이틀곡 《I\'m Your Girl》로 공식 데뷔. '
     '데뷔와 동시에 "요정 신드롬"을 일으키며 걸그룹 전성시대의 서막을 열었다.'),
    ('② 전성기 — 정규 1·2집 (1997~1999)',
     '1집 《S.E.S.》로 신드롬을 일으킨 데 이어, 2집 《S.E.S.2》에서 《Dreams Come True》, '
     '《너를 사랑해》 등을 히트시키며 승승장구했다. 발표하는 노래마다 소녀들의 감성을 가득 담아 '
     '남녀노소 폭넓은 지지를 얻었다.'),
    ('③ 음악적 성숙 — 정규 3·4집 (1999~2001)',
     '3집 《Love》에서 《Love》, 《Twilight Zone》 등을 선보이며 음악적 완성도를 높였다. '
     '4집 《A Letter from Greenland》에서는 더욱 성숙한 음악 세계를 펼쳐 평론가들로부터 '
     '"이미지만 아니라 음악까지 인정할 수 있는 최초의 아이돌"이라는 평가를 받았다.'),
    ('④ 마지막 활동 및 해체 (2002~2003)',
     '5집 《Choose My Life-U》를 마지막 정규앨범으로 활동한 뒤 2003년 공식 해체. '
     '싱글 《Remixed》와 스페셜 앨범 《FRIEND》를 끝으로 그룹 활동을 마무리했다.'),
    ('⑤ 20주년 재결성 (2016~2017)',
     '해체 14년 만인 2016~2017년, 걸그룹 최초의 데뷔 20주년 기념 재결성을 이루어냈다. '
     'SM STATION을 통해 《Love Story》를 발매하고, 스페셜 앨범 《REMEMBER》를 내놓으며 '
     '단독 콘서트와 팬미팅을 진행해 팬들과 다시 만났다.'),
]

for title, content in history:
    add_para(doc, title, bold=True)
    add_para(doc, content, indent=True)
    doc.add_paragraph()

# ── 4. 해외 활동 ──────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '4. 해외 활동', level=2)

add_para(doc, '■ 일본 활동', bold=True)
add_para(doc,
    'SM엔터테인먼트의 일본 시장 진출 선구자로서 일본에서도 활동했다. 일본 현지 음반 발매 및 '
    '공연을 통해 한류의 초석을 다지는 역할을 했으나, 언어 장벽과 현지 시장 특성으로 인해 '
    '한국에서만큼의 폭발적인 성공을 거두지는 못했다.',
    indent=True
)
add_para(doc, '■ 대만 활동', bold=True)
add_para(doc,
    '1999년~2001년 총 5차례 대만을 방문하며 아시아 시장으로 활동 영역을 넓혔다. '
    '걸그룹 최초로 해외 공식 팬클럽이 대만에서 탄생했을 만큼 현지에서도 상당한 인기를 얻었으며, '
    'ICRT 지진 자선공연 참석 등 사회적 활동도 병행했다.',
    indent=True
)
doc.add_paragraph()

# ── 5. 주요 음반 ──────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '5. 주요 음반', level=2)

albums = [
    '정규 1집 《S.E.S.》(1997) — 《I\'m Your Girl》, 《Oh My Love》',
    '정규 2집 《S.E.S.2》(1998) — 《Dreams Come True》, 《너를 사랑해》',
    '정규 3집 《Love》(1999) — 《Love》, 《Twilight Zone》, 《샤랄라》',
    '정규 4집 《A Letter from Greenland》(2001)',
    '스페셜 4.5집 《Surprise》(2001) — 일본 활동 집대성',
    '정규 5집 《Choose My Life-U》(2002) — 마지막 정규앨범',
    '스페셜 앨범 《REMEMBER》(2017) — 20주년 재결성 기념',
]
for a in albums:
    add_bullet(doc, a)
doc.add_paragraph()

# ── 6. 평가 및 영향력 ────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '6. 평가 및 걸그룹 영향력', level=2)

add_para(doc,
    'S.E.S.는 대한민국 가요계에서 걸그룹이라는 개념을 대중에게 처음으로 각인시킨 그룹이다. '
    '이들의 성공 이후 수많은 3인조 걸그룹이 등장했으며, K-Pop 걸그룹의 원형(原型)으로 평가받는다.'
)
doc.add_paragraph()

evaluations = [
    '음악 평론가 김봉현: "이미지만 아니라 음악까지 인정할 수 있는 최초의 아이돌"',
    '음악 평론가 강명석: 《I\'m Your Girl》은 1절부터 대중을 사로잡았으며, 《Love》·《Be Natural》은 걸그룹 음악의 새 지평을 열었다.',
    '대중성과 음악성을 동시에 인정받은 1990년대 대표 걸그룹',
    '"누나부대"라는 새로운 팬덤 문화를 창출 — 성별·세대를 넘는 팬층 형성',
    '이후 등장하는 수많은 K-Pop 걸그룹의 롤모델로 지속적으로 언급됨',
]
for e in evaluations:
    add_bullet(doc, e)
doc.add_paragraph()

add_divider(doc)
doc.add_paragraph()

# ── 결론 ─────────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '결론', level=2)

add_para(doc,
    'S.E.S.는 1997년 데뷔 이후 약 6년간 한국 가요계를 이끈 전설적인 걸그룹이다. '
    '청순한 이미지와 완성도 높은 음악으로 대중성과 예술성을 동시에 인정받았으며, '
    '"요정 신드롬"이라는 신조어를 탄생시킬 만큼 당대 최고의 인기를 누렸다. '
    '2016~2017년 20주년 재결성을 통해 그 인기가 현재도 유효함을 증명했으며, '
    'K-Pop 걸그룹의 역사에서 빼놓을 수 없는 아이콘으로 영원히 기억될 것이다.'
)

# ── 저장 ─────────────────────────────────────────────
output_path = r'C:\Users\SBS\Downloads\agent ai-2\pw\SES_요약보고서.docx'
doc.save(output_path)
print('saved:', output_path)
