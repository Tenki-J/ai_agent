# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

FONT_NAME = '맑은 고딕'

doc = Document()

# 문서 기본 스타일(Normal)에 동아시아 폰트 설정
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)

# 페이지 여백
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)


def apply_font(run, name=FONT_NAME):
    """run 에 한글 폰트를 안전하게 적용"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)


def set_heading(paragraph, text, level=2, color=RGBColor(0x1a, 0x53, 0x76)):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    apply_font(run)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after  = Pt(4)


def add_para(doc, text, size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    apply_font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    apply_font(run)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 표지 ──────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('인공지능(AI) 요약 보고서')
title_run.bold = True
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor(0x1a, 0x53, 0x76)
apply_font(title_run)

doc.add_paragraph()

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run('출처: 위키백과 한국어판 (ko.wikipedia.org/wiki/인공지능)')
sub_run.font.size = Pt(10)
sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
apply_font(sub_run)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('작성일: ' + datetime.date.today().strftime('%Y년 %m월 %d일'))
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
apply_font(date_run)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ── 1. 개요 ───────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '1. 개요 (Definition)', level=2)

add_para(doc,
    '인공지능(人工智能, Artificial Intelligence, AI)은 인간의 학습능력·추론능력·지각능력을 인공적으로 '
    '구현하려는 컴퓨터 과학의 세부 분야이자, 정보공학의 핵심 인프라 기술이다. '
    '인간을 포함한 동물이 지닌 자연 지능(natural intelligence)과는 구별되는 개념으로, '
    '인간의 지능을 모방한 기능을 갖춘 컴퓨터 시스템 및 그 구현 방법론·가능성을 연구하는 '
    '과학기술 분야 전반을 포괄한다.'
)
doc.add_paragraph()

# ── 2. 강AI vs 약AI ──────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '2. 강인공지능 vs 약인공지능', level=2)

add_para(doc, '■ 약인공지능 (Weak AI)', bold=True)
add_para(doc,
    '특정 문제 해결에 특화된 AI로, 사진 속 물체 인식·음성 인식처럼 인간이 쉽게 하지만 '
    '컴퓨터에는 어려웠던 작업을 수행한다. 일반 지능을 목표로 하지 않으며, 실용적 도구로 활용된다.',
    indent=True
)
add_para(doc, '■ 강인공지능 (Strong AI / AGI)', bold=True)
add_para(doc,
    '인간처럼 실제로 사고하여 문제를 해결하는 "일반 지능"을 인공적으로 구현하려는 시도. '
    'AGI 구현에는 추론·문제해결·지식표현·계획수립·학습 등 지능적 특성의 통합이 필요하며, '
    '새로운 환경에서 스스로 학습·적응하는 일반화 능력이 핵심이다.',
    indent=True
)
doc.add_paragraph()

# ── 3. 역사 ───────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '3. 인공지능의 역사', level=2)

history = [
    ('① 태동기 (1943~1956)',
     '수학·철학·공학 등 다양한 분야에서 인공두뇌 가능성 논의. 1950년 앨런 튜링이 "튜링 테스트" 제안. '
     '1956년 다트머스 컨퍼런스(마빈 민스키·존 매카시·클로드 섀넌 등 주도)에서 "인공지능"이라는 '
     '명칭과 학문 분야가 공식 탄생.'),
    ('② 황금기 (1956~1974)',
     '대수학 문제 풀기·기하 정리 증명·영어 학습 등 놀라운 성과 등장. '
     'DARPA·MIT 등에 대규모 자금 투입. 자연어 처리(ELIZA), 탐색 추리, 마이크로월드 연구 등 발전. '
     '연구자들은 "20년 안에 완전한 지능 기계 탄생" 등 과도한 낙관론 표출.'),
    ('③ 제1차 AI 겨울 (1974~1980)',
     '컴퓨터 성능 한계·조합 폭발 문제·상식 지식 부재 등으로 한계 노출. '
     '영국 라이트힐 보고서(1973)로 연구소 해체, DARPA 자금 중단. 뉴럴 네트워크 연구도 일시 중단.'),
    ('④ AI 붐 (1980~1987)',
     '전문가 시스템(XCON 등) 상용화로 재도약. 일본 5세대 컴퓨터 프로젝트(8억 5천만 달러 투자). '
     '1982년 홉필드 신경망·역전파 알고리즘 복원으로 신경망 연구 재개.'),
    ('⑤ 제2차 AI 겨울 (1987~1993)',
     '전문가 시스템의 높은 유지비용·취약성 노출. Lisp 전용 하드웨어 시장 붕괴. '
     '일본 5세대 프로젝트 목표 미달. "AI winter" 용어 등장.'),
    ('⑥ 현대 AI (1993~현재)',
     '1997년 딥블루가 체스 세계 챔피언 카스파로프 격파. 2005년 DARPA 챌린지 자율주행 우승. '
     '2011년 IBM 왓슨 Jeopardy! 우승. 지능형 에이전트 패러다임 정립. '
     '베이지안 네트워크·딥러닝 등 수학적 기법 고도화. AI가 산업 전반에 실용적으로 통합.'),
]

for title, content in history:
    add_para(doc, title, bold=True)
    add_para(doc, content, indent=True)
    doc.add_paragraph()

# ── 4. 주요 연구 분야 ────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '4. 주요 연구 분야 및 접근 방식', level=2)

fields = [
    '탐색 추리(Search Reasoning): 목표 달성을 위한 단계적 경로 탐색 알고리즘',
    '자연어 처리(NLP): 인간 언어로 컴퓨터와 소통 (ELIZA → 현대 LLM)',
    '신경망 / 딥러닝: 뇌의 신경 구조를 모방한 학습 모델',
    '전문가 시스템: 특정 도메인 지식 기반의 규칙 추론 시스템',
    '컴퓨터 비전: 이미지·영상 인식 및 해석',
    '로보틱스: 물리적 환경을 인식·이동·조작하는 지능형 로봇',
    '지능형 에이전트: 환경 인식 후 목표 극대화 행동을 취하는 자율 시스템',
]
for f in fields:
    add_bullet(doc, f)
doc.add_paragraph()

# ── 5. 한계 ───────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '5. 인공지능의 한계 및 도전 과제', level=2)

limits = [
    '컴퓨터 연산 능력의 한계 (초기): 메모리·처리 속도 부족으로 복잡한 문제 해결 불가',
    '조합 폭발 문제: 경우의 수가 기하급수적으로 증가해 최적해 탐색 비용이 과도함',
    '상식 지식 부재: 방대한 일상 지식의 데이터베이스화 및 자동 학습이 어려움',
    '모라벡의 패러독스: 논리 문제는 쉽지만 얼굴 인식·보행 등 감각운동 과제는 어려움',
    '프레임 문제: 행동 후 변화한 세계 상태를 효율적으로 표현·추론하기 어려움',
    'AI 보안 위협: Poisoning Attack, Evasion Attack, Model Extraction, Inversion Attack 등',
]
for l in limits:
    add_bullet(doc, l)
doc.add_paragraph()

# ── 6. 실용적 응용 ───────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '6. 실용적 응용 분야', level=2)

apps = [
    '의료: 질환 진단 (MYCIN 등 전문가 시스템, 현대 의료 영상 분석)',
    '게임: 체스(딥블루), 바둑(알파고), 퀴즈(왓슨)',
    '자율주행: DARPA 챌린지 우승 자율주행 차량',
    '자연어: 음성 인식, 기계 번역, 챗봇, 대형 언어 모델(LLM)',
    '산업: 산업 로봇공학, 데이터 마이닝, 은행 소프트웨어, 광학 문자 인식(OCR)',
    '추천 시스템: 검색 엔진, 콘텐츠 추천 알고리즘',
]
for a in apps:
    add_bullet(doc, a)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ── 결론 ─────────────────────────────────────────────
h = doc.add_heading('', level=2)
set_heading(h, '결론', level=2)

add_para(doc,
    '인공지능은 1956년 학문으로 공식 탄생한 이후 두 차례의 "AI 겨울"을 거치며 '
    '부침을 반복했으나, 컴퓨터 성능 향상·빅데이터·딥러닝 등의 결합으로 현재 전례 없는 '
    '발전을 이루고 있다. 약인공지능은 이미 산업 전반에 깊숙이 통합되었으며, '
    '강인공지능(AGI) 실현을 향한 연구도 지속되고 있다. '
    '동시에 AI 윤리·보안·사회적 영향에 대한 논의도 함께 심화되고 있다.'
)

# ── 저장 ─────────────────────────────────────────────
output_path = r'C:\Users\SBS\Downloads\agent ai-2\pw\인공지능_요약보고서.docx'
doc.save(output_path)
print('saved:', output_path)
