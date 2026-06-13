"""
Landing Page Comparison Report Generator
Creates a Word document comparing 3 landing page types:
1. SaaS Product
2. App Download
3. Email Collection
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\SBS\.claude\skills\docx-creator-workspace\iteration-1\eval-landing-comparison\without_skill\outputs\landing_page_comparison.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def heading(doc, text, level=1, color="1F3864"):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    )
    return p


def body(doc, text, bold=False, italic=False, indent=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def bullet(doc, text, level=0):
    """Add a bullet-list item."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_header_row(table, texts, bg="1F3864", fg="FFFFFF"):
    """Style the first row as a header row."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(
            int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
        )


def fill_cell(cell, text, bold=False, center=False, bg=None, font_size=10.5):
    """Fill a table cell with formatted text."""
    if bg:
        set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


# ── Data ─────────────────────────────────────────────────────────────────────

TYPES = ["SaaS 제품 (SaaS Product)", "앱 다운로드 (App Download)", "이메일 수집 (Email Collection)"]

HERO_DATA = {
    "SaaS 제품 (SaaS Product)": {
        "headline": "고객 전환율을 2배 높이는 CRM 솔루션",
        "subheadline": "영업팀을 위한 올인원 고객 관리 플랫폼 — 14일 무료 체험",
        "visual": "제품 대시보드 스크린샷 또는 인터랙티브 데모 GIF",
        "social_proof": "별점 4.8 / 가입 기업 500+ / 주요 고객사 로고",
        "cta_primary": "무료로 시작하기",
        "cta_secondary": "제품 데모 보기",
        "bg_style": "밝은 배경, 제품 UI 강조, 깔끔한 화이트 or 라이트 그레이",
    },
    "앱 다운로드 (App Download)": {
        "headline": "이동 중에도 끊김 없는 업무 관리",
        "subheadline": "iOS · Android 무료 다운로드 — 설치 후 3분이면 시작",
        "visual": "모바일 기기 목업(Mockup)에 앱 스크린샷 삽입",
        "social_proof": "App Store 4.9★ · 다운로드 100만+ · 에디터 추천",
        "cta_primary": "App Store에서 다운로드",
        "cta_secondary": "Google Play에서 다운로드",
        "bg_style": "그라디언트 배경, 모바일 목업 크게 배치, 스토어 배지 강조",
    },
    "이메일 수집 (Email Collection)": {
        "headline": "매주 월요일, 성장하는 스타트업의 인사이트를 받아보세요",
        "subheadline": "구독자 12,000명이 먼저 읽는 뉴스레터 — 무료 · 언제든 해지 가능",
        "visual": "뉴스레터 미리보기 이미지 또는 혜택 아이콘 그리드",
        "social_proof": "오픈율 45% · 구독자 12,000+ · 에디터 추천 2회 선정",
        "cta_primary": "지금 무료 구독하기",
        "cta_secondary": "지난 호 미리보기",
        "bg_style": "단색 또는 일러스트 배경, 이메일 입력 폼 중앙 배치",
    },
}

CTA_DATA = {
    "SaaS 제품 (SaaS Product)": [
        "무료로 시작하기 (Start for Free)",
        "14일 무료 체험 — 신용카드 불필요",
        "지금 바로 데모 신청하기",
        "팀 플랜 견적 받기",
        "가격 비교 보기",
    ],
    "앱 다운로드 (App Download)": [
        "App Store에서 다운로드",
        "Google Play에서 받기",
        "지금 무료로 설치하기",
        "3분 만에 시작하세요",
        "한 번 써보면 못 돌아갑니다 →",
    ],
    "이메일 수집 (Email Collection)": [
        "무료로 구독하기",
        "이메일을 입력하고 첫 호를 받아보세요",
        "지금 가입 — 언제든 해지 가능",
        "12,000명과 함께 성장하세요",
        "첫 번째 이슈 바로 받기 →",
    ],
}

SECTIONS_DATA = {
    "SaaS 제품 (SaaS Product)": [
        ("Hero Section", "헤드라인, 서브카피, 제품 스크린샷, CTA 버튼 2개"),
        ("Social Proof", "로고 배너(고객사), 수치 강조(가입자 수, 평점)"),
        ("Feature Highlights", "핵심 기능 3~5가지 — 아이콘 + 짧은 설명"),
        ("How It Works", "온보딩 3단계 프로세스 (설치 → 설정 → 결과)"),
        ("Pricing Table", "무료/기본/프로 플랜 비교표"),
        ("Testimonials", "실제 고객 인용문 + 직함/회사"),
        ("FAQ", "자주 묻는 질문 아코디언"),
        ("Final CTA", "페이지 하단 재강조 CTA + 신뢰 배지"),
    ],
    "앱 다운로드 (App Download)": [
        ("Hero Section", "앱 목업, 헤드라인, 스토어 배지 CTA"),
        ("App Preview", "주요 화면 스크린샷 슬라이더 or 캐러셀"),
        ("Key Features", "앱의 핵심 기능 3~4가지 — 모바일 친화적 레이아웃"),
        ("Social Proof", "스토어 평점, 리뷰 수, 수상 배지"),
        ("User Reviews", "실제 앱스토어 리뷰 발췌"),
        ("Platform Info", "지원 OS 버전, 용량, 호환 기기"),
        ("Download CTA", "iOS / Android 배지 + QR코드"),
        ("Privacy Note", "데이터 정책 요약 (짧게)"),
    ],
    "이메일 수집 (Email Collection)": [
        ("Hero Section", "헤드라인, 가치 제안, 이메일 입력 폼 + CTA 버튼"),
        ("Value Proposition", "구독 시 얻을 수 있는 혜택 3가지 불릿"),
        ("Sample Content", "지난 호 하이라이트 또는 콘텐츠 미리보기"),
        ("Social Proof", "구독자 수, 오픈율, 추천 미디어"),
        ("Testimonials", "기존 구독자 인용문"),
        ("Frequency & Privacy", "발송 주기, 해지 방법, 스팸 없음 보장"),
        ("Secondary CTA", "페이지 중간/하단 반복 폼"),
        ("Footer", "SNS 링크, 지난 호 아카이브"),
    ],
}

# Summary comparison table data
COMPARISON_ROWS = [
    ("주요 목표", "제품 체험 / 유료 전환", "앱 설치 유도", "이메일 주소 확보"),
    ("Hero 핵심 요소", "제품 UI + 가치 수치", "모바일 목업 + 스토어 배지", "이메일 폼 중앙 배치"),
    ("주요 CTA", "무료 체험 / 데모 신청", "앱 다운로드 (iOS/Android)", "무료 구독하기"),
    ("Social Proof 유형", "고객사 로고 + 수치", "스토어 평점 + 리뷰 수", "구독자 수 + 오픈율"),
    ("필수 섹션 수", "8개", "8개", "8개"),
    ("신뢰 구축 방식", "Testimonials + 가격 비교", "리뷰 + OS 호환 정보", "구독자 인용 + 스팸 없음"),
    ("전환 마찰 최소화", "신용카드 불필요 강조", "QR코드 제공", "원클릭 구독 + 해지 보장"),
    ("콘텐츠 길이", "중간~길게 (기능·가격 필요)", "짧게~중간 (시각 중심)", "짧게 (폼 집중)"),
]


# ── Document Builder ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("랜딩페이지 유형별 비교 보고서")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Landing Page Type Comparison Report")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    sub_run.italic = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("2026년 6월 7일")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

    doc.add_paragraph()

    # ── Introduction ──────────────────────────────────────────────────────────
    heading(doc, "1. 개요 (Overview)", level=1)
    body(doc,
         "본 보고서는 디지털 마케팅에서 자주 활용되는 세 가지 랜딩페이지 유형을 체계적으로 비교·분석합니다. "
         "각 유형의 목적, Hero 섹션 구성 방식, CTA 문구 전략, 그리고 필수 섹션 목록을 정리하고, "
         "마지막에 통합 비교표로 요약합니다.")
    body(doc, "분석 대상 유형:")
    for t in TYPES:
        bullet(doc, t)
    doc.add_paragraph()

    # ── Per-type sections ─────────────────────────────────────────────────────
    type_colors = ["2E74B5", "538135", "7030A0"]  # blue, green, purple

    for idx, lp_type in enumerate(TYPES):
        color = type_colors[idx]
        short_name = lp_type.split("(")[0].strip()

        # Section heading
        heading(doc, f"{idx + 2}. {lp_type}", level=1, color=color)

        # ── Hero 구성 ──
        heading(doc, "2-1. Hero 섹션 구성" if idx == 0 else f"{idx + 2}-1. Hero 섹션 구성",
                level=2, color=color)

        hero = HERO_DATA[lp_type]
        # Hero table
        hero_table = doc.add_table(rows=8, cols=2)
        hero_table.style = "Table Grid"
        set_cell_borders(hero_table)

        hero_rows = [
            ("헤드라인 (Headline)", hero["headline"]),
            ("서브헤드라인 (Sub-headline)", hero["subheadline"]),
            ("비주얼 요소 (Visual)", hero["visual"]),
            ("소셜 증명 (Social Proof)", hero["social_proof"]),
            ("Primary CTA", hero["cta_primary"]),
            ("Secondary CTA", hero["cta_secondary"]),
            ("배경 스타일 (Background)", hero["bg_style"]),
        ]

        # Header row
        add_header_row(hero_table, ["항목", "내용"], bg=color)

        # Fill rows
        alt_colors = ["F2F7FF", "FFFFFF"]
        for r_idx, (label, value) in enumerate(hero_rows):
            row = hero_table.rows[r_idx + 1]
            fill_cell(row.cells[0], label, bold=True,
                      bg=alt_colors[r_idx % 2], font_size=10)
            fill_cell(row.cells[1], value,
                      bg=alt_colors[r_idx % 2], font_size=10)

        # Column widths
        for row in hero_table.rows:
            row.cells[0].width = Cm(4.5)
            row.cells[1].width = Cm(11.5)

        doc.add_paragraph()

        # ── CTA 문구 예시 ──
        heading(doc,
                "2-2. CTA 문구 예시" if idx == 0 else f"{idx + 2}-2. CTA 문구 예시",
                level=2, color=color)
        body(doc, f"{short_name} 랜딩페이지에서 효과적인 CTA 문구 예시:")
        for cta in CTA_DATA[lp_type]:
            bullet(doc, cta)

        doc.add_paragraph()

        # ── 필수 섹션 목록 ──
        heading(doc,
                "2-3. 필수 섹션 목록" if idx == 0 else f"{idx + 2}-3. 필수 섹션 목록",
                level=2, color=color)

        sec_table = doc.add_table(rows=len(SECTIONS_DATA[lp_type]) + 1, cols=3)
        sec_table.style = "Table Grid"
        set_cell_borders(sec_table)

        add_header_row(sec_table, ["#", "섹션명", "주요 구성 요소"], bg=color)

        for s_idx, (sec_name, sec_desc) in enumerate(SECTIONS_DATA[lp_type]):
            row = sec_table.rows[s_idx + 1]
            bg = alt_colors[s_idx % 2]
            fill_cell(row.cells[0], str(s_idx + 1), center=True, bg=bg, font_size=10)
            fill_cell(row.cells[1], sec_name, bold=True, bg=bg, font_size=10)
            fill_cell(row.cells[2], sec_desc, bg=bg, font_size=10)

        for row in sec_table.rows:
            row.cells[0].width = Cm(1.0)
            row.cells[1].width = Cm(4.5)
            row.cells[2].width = Cm(10.5)

        doc.add_paragraph()

    # ── Summary Comparison Table ──────────────────────────────────────────────
    heading(doc, f"{len(TYPES) + 2}. 전체 비교표 (Summary Comparison)", level=1, color="1F3864")
    body(doc, "세 가지 랜딩페이지 유형의 핵심 요소를 한눈에 비교합니다.")
    doc.add_paragraph()

    num_cols = 4  # 항목 + 3 types
    num_rows = len(COMPARISON_ROWS) + 1  # header + data

    comp_table = doc.add_table(rows=num_rows, cols=num_cols)
    comp_table.style = "Table Grid"
    set_cell_borders(comp_table)

    # Header row
    header_texts = ["비교 항목"] + [t.split("(")[0].strip() for t in TYPES]
    add_header_row(comp_table, header_texts, bg="1F3864")

    col_bg = ["FFF2CC", "E2EFDA", "EAD1FF"]  # yellow, green, purple tint

    for r_idx, (label, saas_val, app_val, email_val) in enumerate(COMPARISON_ROWS):
        row = comp_table.rows[r_idx + 1]
        row_bg = "F7F7F7" if r_idx % 2 == 0 else "FFFFFF"
        fill_cell(row.cells[0], label, bold=True, bg=row_bg, font_size=10)
        fill_cell(row.cells[1], saas_val, bg=col_bg[0] if r_idx % 2 == 0 else "FFFDE7", font_size=10)
        fill_cell(row.cells[2], app_val, bg=col_bg[1] if r_idx % 2 == 0 else "F1F8E9", font_size=10)
        fill_cell(row.cells[3], email_val, bg=col_bg[2] if r_idx % 2 == 0 else "F3E5FF", font_size=10)

    # Column widths for comparison table
    col_widths = [Cm(4.0), Cm(4.5), Cm(4.5), Cm(4.5)]
    for row in comp_table.rows:
        for c_idx, width in enumerate(col_widths):
            row.cells[c_idx].width = width

    doc.add_paragraph()

    # ── Conclusion ────────────────────────────────────────────────────────────
    heading(doc, f"{len(TYPES) + 3}. 결론 및 선택 가이드", level=1)
    body(doc, "랜딩페이지 유형 선택 시 다음 기준을 참고하십시오.", bold=True)
    doc.add_paragraph()

    guide_items = [
        ("SaaS 제품 랜딩페이지", "복잡한 기능 설명, 가격 비교, 신뢰 구축이 필요한 B2B/B2C SaaS 서비스에 적합합니다. "
         "전환 주기가 길기 때문에 무료 체험 CTA로 진입 장벽을 낮추는 것이 핵심입니다."),
        ("앱 다운로드 랜딩페이지", "모바일 앱의 설치를 직접 유도할 때 사용합니다. 시각적 목업과 스토어 평점이 "
         "신뢰를 만들며, QR코드로 마찰을 최소화합니다."),
        ("이메일 수집 랜딩페이지", "리드를 확보하거나 뉴스레터 구독자를 늘릴 때 가장 효율적입니다. "
         "폼이 단순할수록 전환율이 높으며, 혜택 명시와 해지 보장이 필수입니다."),
    ]

    for title_text, desc_text in guide_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run_title = p.add_run(f"{title_text}: ")
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_desc = p.add_run(desc_text)
        run_desc.font.size = Pt(11)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("— 백곰이 자동 생성 보고서 | 2026-06-07 —")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    footer_run.italic = True

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"[OK] Document saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    size_kb = os.path.getsize(path) / 1024
    print(f"[OK] File size: {size_kb:.1f} KB")

    # Report structure
    print("\n=== Document Structure ===")
    print("1. 개요 (Overview)")
    print("2. SaaS 제품 랜딩페이지")
    print("   2-1. Hero 섹션 구성 (table: 7 rows)")
    print("   2-2. CTA 문구 예시 (5 bullet items)")
    print("   2-3. 필수 섹션 목록 (table: 8 sections)")
    print("3. 앱 다운로드 랜딩페이지")
    print("   3-1. Hero 섹션 구성 (table: 7 rows)")
    print("   3-2. CTA 문구 예시 (5 bullet items)")
    print("   3-3. 필수 섹션 목록 (table: 8 sections)")
    print("4. 이메일 수집 랜딩페이지")
    print("   4-1. Hero 섹션 구성 (table: 7 rows)")
    print("   4-2. CTA 문구 예시 (5 bullet items)")
    print("   4-3. 필수 섹션 목록 (table: 8 sections)")
    print("5. 전체 비교표 (Summary Comparison table: 8 comparison rows x 4 cols)")
    print("6. 결론 및 선택 가이드")
